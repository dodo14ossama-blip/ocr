from flask import Flask, request, jsonify
from flask_cors import CORS
import re
import random
import io
import json
from datetime import datetime

app = Flask(__name__)
CORS(app)

# ==================== Swagger UI HTML ====================
SWAGGER_UI_HTML = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>🧬 Medical Data Extractor API - Swagger UI</title>
    <link rel="stylesheet" type="text/css" href="https://cdn.jsdelivr.net/npm/swagger-ui-dist@5/swagger-ui.css">
    <link rel="icon" type="image/png" href="https://cdn.jsdelivr.net/npm/swagger-ui-dist@5/favicon-32x32.png">
    <style>
        body { margin: 0; padding: 0; }
        .topbar { background-color: #1a237e; padding: 10px; }
        .topbar .wrapper { max-width: 1460px; margin: 0 auto; padding: 0 20px; }
        .topbar a { color: white; font-size: 1.5em; text-decoration: none; font-weight: bold; }
        .topbar a span { color: #64b5f6; }
    </style>
</head>
<body>
    <div class="topbar">
        <div class="wrapper">
            <a href="/">🧬 Medical Data Extractor <span>API</span></a>
            <a href="/" style="font-size: 0.8em; margin-left: 20px;">🏠 Home</a>
            <a href="/swagger" style="font-size: 0.8em; margin-left: 20px;">📚 Swagger UI</a>
        </div>
    </div>
    <div id="swagger-ui"></div>
    <script src="https://cdn.jsdelivr.net/npm/swagger-ui-dist@5/swagger-ui-bundle.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/swagger-ui-dist@5/swagger-ui-standalone-preset.js"></script>
    <script>
        window.onload = () => {
            window.ui = SwaggerUIBundle({
                url: "/swagger.json",
                dom_id: '#swagger-ui',
                presets: [SwaggerUIBundle.presets.apis, SwaggerUIStandalonePreset],
                layout: "BaseLayout",
                deepLinking: true,
                tryItOutEnabled: true,
                docExpansion: "list",
                filter: true,
                persistAuthorization: false,
            });
        };
    </script>
</body>
</html>
"""

# ==================== Swagger JSON ====================
@app.route('/swagger.json')
def swagger_json():
    return jsonify({
        "openapi": "3.0.0",
        "info": {
            "title": "🧬 Medical Data Extractor API",
            "description": """
## استخراج البيانات الطبية من أي ملف

### 📌 المميزات:
- 🩺 استخراج العمر، السكر، ضغط الدم، LDL، HDL
- 🧬 حساب نسبة المخاطر الوراثية
- 📁 دعم جميع أنواع الملفات (صور، PDF، Excel، Word، نص)

### 📂 أنواع الملفات المدعومة:
| النوع | الامتدادات |
|-------|-------------|
| 📸 الصور | jpg, png, jpeg, bmp, gif |
| 📄 المستندات | pdf, docx, txt |
| 📊 الجداول | xlsx, xls |

### 🚀 مثال الاستخدام:
```bash
# استخراج البيانات
curl -X POST https://your-api.vercel.app/extract -F "file=@report.pdf"

# حساب المخاطر
curl -X POST https://your-api.vercel.app/predict -F "file=@patient.jpg"
            """,
"version": "3.0.0",
"contact": {
"name": "Medical Data Extractor",
"email": "support@medical-extractor.com"
}
},
"servers": [
{"url": "/", "description": "Current Server"},
{"url": "https://medical-data-extractor.vercel.app", "description": "Vercel Production"}
],
"tags": [
{"name": "Information", "description": "📋 معلومات الـ API"},
{"name": "Extraction", "description": "🔍 استخراج البيانات الطبية"},
{"name": "Prediction", "description": "🎯 توقع المخاطر"}
],
"paths": {
"/": {
"get": {
"tags": ["Information"],
"summary": "🏠 الصفحة الرئيسية",
"description": "يعرض واجهة رفع الملفات",
"responses": {"200": {"description": "HTML page"}}
}
},
"/health": {
"get": {
"tags": ["Information"],
"summary": "❤️ فحص الصحة",
"description": "يتحقق من أن الـ API يعمل بشكل صحيح",
"responses": {
"200": {
"description": "الـ API يعمل",
"content": {
"application/json": {
"example": {"status": "healthy", "timestamp": "2024-01-01T00:00:00"}
}
}
}
}
}
},
"/swagger": {
"get": {
"tags": ["Information"],
"summary": "📚 واجهة Swagger UI",
"description": "يفتح واجهة توثيق الـ API التفاعلية",
"responses": {"200": {"description": "Swagger UI interface"}}
}
},
"/extract": {
"post": {
"tags": ["Extraction"],
"summary": "📊 استخراج البيانات الطبية",
"description": "يرفع ملف ويستخرج منه البيانات الطبية (العمر، السكر، ضغط الدم، LDL، تحاليل الدم)",
"operationId": "extractMedicalData",
"requestBody": {
"required": True,
"content": {
"multipart/form-data": {
"schema": {
"type": "object",
"properties": {
"file": {
"type": "string",
"format": "binary",
"description": "الملف الطبي (صورة، PDF، Excel، Word، نص)",
"example": "patient_report.pdf"
}
},
"required": ["file"]
}
}
}
},
"responses": {
"200": {
"description": "تم استخراج البيانات بنجاح",
"content": {
"application/json": {
"example": {
"success": True,
"filename": "report.pdf",
"extracted_data": {
"age": 58,
"glucose": 165,
"systolic_bp": 145,
"diastolic_bp": 90,
"ldl": 170,
"hemoglobin": 16.5,
"platelets": 242,
"wbc": 4.8,
"genetic_risk_score": 0.7,
"gender": "Male"
}
}
}
}
},
"400": {"description": "لم يتم رفع أي ملف"},
"500": {"description": "خطأ في السيرفر"}
}
}
},
"/predict": {
"post": {
"tags": ["Prediction"],
"summary": "🎯 استخراج البيانات + حساب المخاطر",
"description": "يرفع ملف ويستخرج البيانات الطبية ويحسب نسبة المخاطر الوراثية",
"operationId": "predictRisk",
"requestBody": {
"required": True,
"content": {
"multipart/form-data": {
"schema": {
"type": "object",
"properties": {
"file": {
"type": "string",
"format": "binary",
"description": "الملف الطبي"
}
},
"required": ["file"]
}
}
}
},
"responses": {
"200": {
"description": "تم حساب المخاطر بنجاح",
"content": {
"application/json": {
"example": {
"success": True,
"extracted_data": {
"age": 58,
"glucose": 165,
"systolic_bp": 145,
"diastolic_bp": 90
},
"risk_assessment": {
"score": 0.575,
"percentage": "57.5%",
"category": "Medium Risk 🟡",
"recommendations": ["Monitor health", "Consider genetic counseling", "Improve lifestyle"]
}
}
}
}
}
}
}
}
},
"components": {
"schemas": {
"MedicalData": {
"type": "object",
"properties": {
"age": {"type": "integer", "description": "العمر بالسنوات"},
"glucose": {"type": "number", "description": "نسبة السكر في الدم (mg/dL)"},
"systolic_bp": {"type": "integer", "description": "الضغط الانقباضي (mmHg)"},
"diastolic_bp": {"type": "integer", "description": "الضغط الانبساطي (mmHg)"},
"ldl": {"type": "number", "description": "الكوليسترول الضار LDL (mg/dL)"},
"hemoglobin": {"type": "number", "description": "الهيموجلوبين (g/dL)"},
"platelets": {"type": "number", "description": "الصفائح الدموية (thousands/mm³)"},
"wbc": {"type": "number", "description": "كريات الدم البيضاء (thousands/mm³)"},
"genetic_risk_score": {"type": "number", "description": "نسبة المخاطر الوراثية (0-1)"}
}
},
"RiskAssessment": {
"type": "object",
"properties": {
"score": {"type": "number", "description": "نسبة المخاطر (0-1)"},
"percentage": {"type": "string", "description": "النسبة المئوية"},
"category": {"type": "string", "description": "فئة المخاطر"},
"recommendations": {"type": "array", "items": {"type": "string"}}
}
}
}
}
})

==================== الواجهة الأصلية (HTML_PAGE) ====================
HTML_PAGE = """

<!DOCTYPE html><html> <head> <title>Medical Data Extractor - OCR Support</title> <meta charset="UTF-8"> <style> body { font-family: Arial; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; padding: 20px; } .container { max-width: 900px; margin: 0 auto; background: white; border-radius: 20px; padding: 30px; } h1 { color: #667eea; text-align: center; } .nav-links { text-align: center; margin-bottom: 20px; } .nav-links a { color: #667eea; margin: 0 15px; text-decoration: none; font-weight: bold; } .upload-area { border: 3px dashed #667eea; border-radius: 15px; padding: 40px; text-align: center; margin-bottom: 20px; } .file-label { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 12px 30px; border-radius: 25px; cursor: pointer; display: inline-block; } button { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border: none; padding: 12px 30px; border-radius: 25px; cursor: pointer; margin: 10px; } .result { margin-top: 20px; padding: 20px; background: #f8f9fa; border-radius: 10px; display: none; max-height: 500px; overflow: auto; } .result.show { display: block; } .error { color: red; background: #ffebee; padding: 10px; border-radius: 5px; } pre { white-space: pre-wrap; word-wrap: break-word; background: #fff; padding: 15px; border-radius: 10px; } table { width: 100%; border-collapse: collapse; } th, td { padding: 8px; text-align: left; border-bottom: 1px solid #ddd; } th { background: #667eea; color: white; } </style> </head> <body> <div class="container"> <h1>🧬 Medical Data Extractor</h1> <div class="nav-links"> <a href="/">🏠 Home</a> <a href="/swagger">📚 Swagger UI</a> </div> <p style="text-align:center">Supports: Images (JPG, PNG), PDF, Excel, Word, Text</p> <div class="upload-area"> <form id="uploadForm" enctype="multipart/form-data"> <label for="fileInput" class="file-label">📁 Choose File (Image, PDF, etc.)</label> <input type="file" name="file" id="fileInput" accept=".txt,.pdf,.jpg,.png,.jpeg,.xlsx,.xls,.docx"> <div id="fileName" style="margin-top:10px; color:#666">No file selected</div> <button type="submit">🚀 Extract Data</button> <button type="button" onclick="sendToPredict()">🎯 Predict Risk</button> </form> </div> <div id="result" class="result"></div> </div> <script> document.getElementById('fileInput').onchange = function() { document.getElementById('fileName').innerHTML = this.files[0] ? this.files[0].name : 'No file selected'; }; async function sendRequest(endpoint) { const file = document.getElementById('fileInput').files[0]; if (!file) { alert('Please select a file first'); return; } const resultDiv = document.getElementById('result'); resultDiv.innerHTML = '<div style="text-align:center">🔄 Processing...</div>'; resultDiv.classList.add('show'); const formData = new FormData(); formData.append('file', file); try { const response = await fetch(endpoint, { method: 'POST', body: formData }); const data = await response.json(); if (data.success) { let html = '<h3>✅ Extracted Data:</h3>'; html += '<table><thead><tr>'; Object.keys(data.extracted_data).forEach(k => html += `<th>${k}</th>`); html += '</tr></thead><tbody></tr>'; Object.values(data.extracted_data).forEach(v => html += `<td>${v !== null ? v : '-'}</td>`); html += '</table></tbody></table>'; if (data.risk_assessment) { html += '<h3>🎯 Risk Assessment:</h3>'; html += `<p><strong>Score:</strong> ${data.risk_assessment.score}</p>`; html += `<p><strong>Percentage:</strong> ${data.risk_assessment.percentage}</p>`; html += `<p><strong>Category:</strong> ${data.risk_assessment.category}</p>`; html += `<p><strong>Recommendations:</strong> ${data.risk_assessment.recommendations.join(', ')}</p>`; } if (data.ocr_text) { html += '<details><summary>📄 Extracted Text Preview</summary>'; html += `<pre>${data.ocr_text.substring(0, 1000)}${data.ocr_text.length > 1000 ? '...' : ''}</pre>`;
html += '</details>';
}

resultDiv.innerHTML = html;
} else {
resultDiv.innerHTML = <div class="error">❌ Error: ${data.error}</div>;
}
} catch(err) {
resultDiv.innerHTML = <div class="error">❌ Error: ${err.message}</div>;
}
}

document.getElementById('uploadForm').onsubmit = (e) => {
e.preventDefault();
sendRequest('/extract');
};

function sendToPredict() {
sendRequest('/predict');
}
</script>

</body> </html> """
==================== استخراج النص من الملفات ====================
def extract_text_from_file(content, filename):
"""استخراج النص من الملف - يدعم OCR للصور"""
ext = filename.split('.')[-1].lower()
text = ""
method = ""

try:
if ext == 'txt':
text = content.decode('utf-8', errors='ignore')
method = "Direct text"

elif ext == 'pdf':
try:
import pdfplumber
with pdfplumber.open(io.BytesIO(content)) as pdf:
for page in pdf.pages:
page_text = page.extract_text()
if page_text:
text += page_text + "\n"
method = f"PDF ({len(pdf.pages)} pages)"
except ImportError:
text = "PDF extraction requires pdfplumber"
method = "PDF (limited)"

elif ext in ['xlsx', 'xls']:
try:
import pandas as pd
df = pd.read_excel(io.BytesIO(content))
text = df.to_string()
method = f"Excel ({df.shape[0]} rows)"
except ImportError:
text = "Excel extraction requires pandas"
method = "Excel (limited)"

elif ext == 'docx':
try:
import docx
doc = docx.Document(io.BytesIO(content))
text = "\n".join([p.text for p in doc.paragraphs])
method = f"Word ({len(doc.paragraphs)} paragraphs)"
except ImportError:
text = "Word extraction requires python-docx"
method = "Word (limited)"

elif ext in ['jpg', 'jpeg', 'png', 'bmp', 'gif', 'tiff']:
try:
from PIL import Image
import pytesseract
img = Image.open(io.BytesIO(content))
text = pytesseract.image_to_string(img, lang='eng')
if not text.strip():
text = pytesseract.image_to_string(img, lang='ara')
method = f"OCR ({img.size[0]}x{img.size[1]} px)"
except ImportError as ie:
text = f"OCR requires PIL and pytesseract"
method = "OCR (libraries missing)"
except Exception as oe:
text = f"OCR processing error"
method = "OCR (error)"

else:
text = content.decode('utf-8', errors='ignore')
method = "Raw text"

except Exception as e:
text = f"Error: {str(e)}"
method = "Error"

return text.strip(), method

def extract_medical_data(text):
"""استخراج القيم الطبية من النص"""
data = {
'age': None, 'glucose': None, 'systolic_bp': None, 'diastolic_bp': None,
'ldl': None, 'hdl': None, 'hemoglobin': None, 'platelets': None,
'wbc': None, 'genetic_risk_score': None, 'gender': None, 'genetic_disease': None
}

if not text:
return data

m = re.search(r'(?:age|عمر|Age)[\s:]*(\d+)', text, re.IGNORECASE)
if m: data['age'] = int(m.group(1))

m = re.search(r'(?:glucose|سكر|Glucose|blood sugar)[\s:]*(\d+(?:.\d+)?)', text, re.IGNORECASE)
if m: data['glucose'] = float(m.group(1))

m = re.search(r'(?:blood pressure|الضغط)[\s:]*(\d+)[\s/-]+(\d+)', text, re.IGNORECASE)
if m:
data['systolic_bp'] = int(m.group(1))
data['diastolic_bp'] = int(m.group(2))

m = re.search(r'(?:ldl|LDL)[\s:]*(\d+(?:.\d+)?)', text, re.IGNORECASE)
if m: data['ldl'] = float(m.group(1))

m = re.search(r'(?:Haemoglobin|Hemoglobin|Hb)[\s:]*(\d+(?:.\d+)?)', text, re.IGNORECASE)
if m: data['hemoglobin'] = float(m.group(1))

m = re.search(r'(?:Platelet Count|PLT)[\s:]*(\d+(?:.\d+)?)', text, re.IGNORECASE)
if m: data['platelets'] = float(m.group(1))

m = re.search(r'(?:Total Leucocytic Count|WBC|Leucocytes)[\s:]*(\d+(?:.\d+)?)', text, re.IGNORECASE)
if m: data['wbc'] = float(m.group(1))

m = re.search(r'(?:genetic risk|الخطر الوراثي)[\s:]*(\d+(?:.\d+)?)', text, re.IGNORECASE)
if m: data['genetic_risk_score'] = float(m.group(1))

if re.search(r'\b(?:male|ذكر|Male|M)\b', text, re.IGNORECASE):
data['gender'] = 'Male'
elif re.search(r'\b(?:female|انثى|Female|F|أنثى)\b', text, re.IGNORECASE):
data['gender'] = 'Female'

m = re.search(r'(?:genetic disease|مرض وراثي|Diagnosis|Comment)[\s:]*([A-Za-z\s]+)', text, re.IGNORECASE)
if m:
disease = m.group(1).strip()
if len(disease) > 3 and disease.lower() not in ['none', 'unknown']:
data['genetic_disease'] = disease[:50]

return data

def calculate_risk(data):
"""حساب نسبة المخاطر"""
risk = 0.0

if data.get('age') and data['age'] > 60: risk += 0.25
elif data.get('age') and data['age'] > 40: risk += 0.125

if data.get('glucose') and data['glucose'] > 200: risk += 0.20
elif data.get('glucose') and data['glucose'] > 140: risk += 0.10

if data.get('systolic_bp') and data['systolic_bp'] > 160: risk += 0.15
elif data.get('systolic_bp') and data['systolic_bp'] > 140: risk += 0.075

if data.get('ldl') and data['ldl'] > 190: risk += 0.15
elif data.get('ldl') and data['ldl'] > 130: risk += 0.075

if data.get('genetic_risk_score'): risk += data['genetic_risk_score'] * 0.15

risk = max(0, min(risk, 0.95))

if risk < 0.3:
cat, rec = "Low Risk 🟢", ["Annual checkup", "Healthy diet", "Regular exercise"]
elif risk < 0.6:
cat, rec = "Medium Risk 🟡", ["Monitor health", "Consider genetic counseling", "Improve lifestyle"]
else:
cat, rec = "High Risk 🔴", ["Consult specialist", "Genetic testing", "Immediate lifestyle changes"]

return {
'score': round(risk, 3),
'percentage': f"{risk*100:.1f}%",
'category': cat,
'recommendations': rec[:3]
}

==================== API Endpoints ====================
@app.route('/', methods=['GET'])
def home():
return HTML_PAGE

@app.route('/swagger', methods=['GET'])
def swagger_ui():
return SWAGGER_UI_HTML

@app.route('/health', methods=['GET'])
def health():
return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()})

@app.route('/extract', methods=['POST'])
def extract():
if 'file' not in request.files:
return jsonify({'success': False, 'error': 'No file uploaded'}), 400

file = request.files['file']
if file.filename == '':
return jsonify({'success': False, 'error': 'No file selected'}), 400

try:
content = file.read()
text, method = extract_text_from_file(content, file.filename)
data = extract_medical_data(text)
data['person_id'] = f"P{random.randint(100000, 999999)}"

return jsonify({
'success': True,
'filename': file.filename,
'extraction_method': method,
'extracted_data': data,
'ocr_text': text[:1500] if text else ""
})
except Exception as e:
return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/predict', methods=['POST'])
def predict():
if 'file' not in request.files:
return jsonify({'success': False, 'error': 'No file uploaded'}), 400

file = request.files['file']
if file.filename == '':
return jsonify({'success': False, 'error': 'No file selected'}), 400

try:
content = file.read()
text, method = extract_text_from_file(content, file.filename)
data = extract_medical_data(text)
risk = calculate_risk(data)
data['person_id'] = f"P{random.randint(100000, 999999)}"

return jsonify({
'success': True,
'filename': file.filename,
'extraction_method': method,
'extracted_data': data,
'risk_assessment': risk
})
except Exception as e:
return jsonify({'success': False, 'error': str(e)}), 500

if name == 'main':
print("=" * 60)
print("🧬 Medical Data Extractor API - with Swagger UI")
print("=" * 60)
print("📍 http://localhost:5000")
print("📚 Swagger UI: http://localhost:5000/swagger")
print("=" * 60)
app.run(host='0.0.0.0', port=5000, debug=True
